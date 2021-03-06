from django.shortcuts import render,get_object_or_404, get_list_or_404, redirect
from django.template.loader import render_to_string, get_template
from django.http import JsonResponse, HttpResponse
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib import messages
from .forms import AuthorityRequestingForm, LocationObjectiveForm, ForensicScanForm, ReportStatusForm, DiscussionConclusionForm, CidNumberForm, ReportForm, NatureOfActionForm, TypeItemForm, TypeItemByNatureOfActionForm, Item2Form
from PericiasMedicas.person.models import ProfilePersonType,Cid10
from .models import AuthorityRequesting, ReportStatus, CidNumber, LocationObjective, ForensicScan, NatureOfAction, DiscussionConclusion, Report, MedicalDocument, TypeItem, TypeItemByNatureOfAction, Item2
from PericiasMedicas.company.models import Department, Company
from PericiasMedicas.person.models import Person, ProfilePersonType, Doctor
from docxtpl import DocxTemplate, RichText # Modulo para exportar em DOCX
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .utils import render_to_pdf #Modulo de exportar PDF
from django.utils.html import strip_tags
from datetime import datetime
from bs4 import BeautifulSoup
from django.db import IntegrityError


############ GET DEPARTMENT #############
#Preciso do ID do Departmento que o usuário pertence           
def get_department(user):
    department = Department.objects.filter(profilepersontype__person__cpf=user)    
    department_id = list()
    for cc in department:        
        department_id.append(cc.id) 
    return department_id

####### FIM GET DEPARTMENT #############

######### AuthorityRequesting ################
@login_required
def authorityrequesting_create(request):
    template_name = 'authorityrequesting/form.html'
    data = {}
    data['title'] = "Cadastro Autoridades Requisitantes"        
    if request.method == 'POST':        
        form = AuthorityRequestingForm(request.POST, department_id=get_department(request.user.username))
        if form.is_valid():
            form.save()                        
            return redirect('url_authorityrequesting_list')
        else:
            print("algo não está valido.")

    form = AuthorityRequestingForm(department_id=get_department(request.user.username))                 
    data['form'] = form
    return render(request,template_name,data)

@login_required
def authorityrequesting_list(request):
    template_name = "authorityrequesting/list.html"
    title = "Lista das Autoridades Requisitantes"
    #authorityrequesting = AuthorityRequesting.objects.all()    
    authorityrequesting = AuthorityRequesting.objects.filter(doctor__profile_person_type__department__in=get_department(request.user.username))
    context = {
        'authorityrequesting': authorityrequesting,
        'title': title      
    }
    return render(request,template_name,context)

@login_required
def authorityrequesting_detail(request, pk=None, *args, **kwargs):
    template_name = "authorityrequesting/detail.html"
    authorityrequesting = get_object_or_404(AuthorityRequesting,pk=pk)
    context = {
        'authorityrequesting': authorityrequesting
    }
    return render(request, template_name, context)

@login_required
def authorityrequesting_edit(request, pk):    
    template_name='authorityrequesting/form.html'
    data = {} 
    authorityrequesting = get_object_or_404(AuthorityRequesting, pk=pk)        
    user_created = authorityrequesting.user_created # Esta linha faz com que o user_created não seja modificado, para mostrar quem criou esta pessoa
    form = AuthorityRequestingForm(request.POST or None, instance=authorityrequesting, department_id=get_department(request.user.username))
    if form.is_valid():        
        authorityrequesting = form.save(commit=False)
        authorityrequesting.user_created = user_created               
        authorityrequesting.save()
        return redirect('url_authorityrequesting_detail',pk=pk)    
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def authorityrequesting_delete(request,pk):
    authorityrequesting = get_object_or_404(AuthorityRequesting, pk=pk)    
    if request.method == 'GET':     
        try:   
            authorityrequesting.delete()
            messages.success(request, 'Ação concluída com sucesso.')
            return redirect('url_authorityrequesting_list')
        except IntegrityError:
            messages.warning(request, 'Não foi possível excluir. Possui dependência.')
            return redirect('url_authorityrequesting_list')    
    else:
        messages.warning(request, 'Ação não concluída.')
        return redirect('url_authorityrequesting_list')

######### FIM AuthorityRequesting ############

######## LOCAL OBJECTIVE ####################

@login_required
def locationobjective_create(request):
    template_name = 'locationobjective/form.html'
    data = {}
    data['title'] = "Cadastro dos Locaise Objetivos"     
    if request.method == 'POST':        
        form = LocationObjectiveForm(request.POST, department_id=get_department(request.user.username))
        if form.is_valid():
            form.save()                        
            return redirect('url_locationobjectives_list')
        else:
            print("algo não está valido.")
    else:
        form = LocationObjectiveForm(department_id=get_department(request.user.username))             
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def locationobjectives_list(request):
    template_name = "locationobjective/list.html"
    title = "Lista - Locais e Objetivos do Exames Periciais"    
    #localobjectives = LocalObjective.objects.all()                    
    locationobjectives = LocationObjective.objects.filter(doctor__profile_person_type__department_id__in=get_department(request.user.username))            
    context = {
        'locationobjectives': locationobjectives,
        'title': title,                      
    }
    return render(request,template_name,context)

@login_required
def locationobjective_detail(request, pk=None, *args, **kwargs):
    template_name = "locationobjective/detail.html"
    locationobjective = get_object_or_404(LocationObjective,pk=pk)
    context = {
        'locationobjective': locationobjective
    }
    return render(request, template_name, context)

@login_required
def locationobjective_edit(request, pk):    
    template_name='locationobjective/form.html'
    data = {} 
    data['title'] = "Cadastro Circunstâncias"         
    locationobjective = get_object_or_404(LocationObjective, pk=pk)        
    user_created = locationobjective.user_created # Esta linha faz com que o user_created não seja modificado, para mostrar quem criou esta pessoa
    form = LocationObjectiveForm(request.POST or None, instance=locationobjective, department_id=get_department(request.user.username))
    if form.is_valid():        
        locationobjective = form.save(commit=False)
        locationobjective.user_created = user_created               
        locationobjective.save()
        return redirect('url_locationobjective_detail',pk=pk)    
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def locationobjective_delete(request,pk):
    locationobjective = get_object_or_404(LocationObjective, pk=pk)    
    if request.method == 'GET':  
        try:       
            locationobjective.delete()
            messages.success(request, 'Ação concluída com sucesso.')
            return redirect('url_locationobjectives_list')
        except IntegrityError:
            messages.warning(request, 'Não foi possível excluir. Possui dependência.')
            return redirect('url_locationobjectives_list')

    else:
        messages.warning(request, 'Ação não concluída.')
        return redirect('url_locationobjectives_list')

######## FIM LOCAL OBJECTIVE ##############

######### ForensicScan #######################

@login_required
def forensicscan_create(request):
    template_name = 'forensicscan/form.html'
    data = {}
    data['title'] = "Cadastro Circunstâncias"     
    if request.method == 'POST':        
        form = ForensicScanForm(request.POST, department_id=get_department(request.user.username))
        if form.is_valid():
            form.save()                        
            return redirect('url_forensicscans_list')
        else:
            print("algo não está valido.",form.errors)
    else:
        form = ForensicScanForm(department_id=get_department(request.user.username))             
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def forensicscans_list(request):
    template_name = "forensicscan/list.html"
    title = "Lista dos Modelos de Anamnese"    
    #forensicscans = ForensicScan.objects.all()                    
    forensicscans = ForensicScan.objects.filter(doctor__profile_person_type__department_id__in=get_department(request.user.username))            
    context = {
        'forensicscans': forensicscans,
        'title': title,                      
    }
    return render(request,template_name,context)

@login_required
def forensicscan_detail(request, pk=None, *args, **kwargs):
    template_name = "forensicscan/detail.html"
    forensicscan = get_object_or_404(ForensicScan,pk=pk)
    context = {
        'forensicscan': forensicscan
    }
    return render(request, template_name, context)

@login_required
def forensicscan_edit(request, pk):    
    template_name='forensicscan/form.html'
    data = {} 
    data['title'] = "Cadastro dos Modelos de Anamnese"         
    forensicscan = get_object_or_404(ForensicScan, pk=pk)        
    user_created = forensicscan.user_created # Esta linha faz com que o user_created não seja modificado, para mostrar quem criou esta pessoa
    form = ForensicScanForm(request.POST or None, instance=forensicscan, department_id=get_department(request.user.username))
    if form.is_valid():        
        forensicscan = form.save(commit=False)
        forensicscan.user_created = user_created               
        forensicscan.save()
        return redirect('url_forensicscan_detail',pk=pk)    
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def forensicscan_delete(request,pk):
    forensicscan = get_object_or_404(ForensicScan, pk=pk)    
    if request.method == 'GET':      
        try:  
            forensicscan.delete()
            messages.success(request, 'Ação concluída com sucesso.')
            return redirect('url_forensicscans_list')
        except IntegrityError:
            messages.warning(request, 'Não foi possível excluir. Possui dependência.')
            return redirect('url_forensicscans_list')    
    else:
        messages.warning(request, 'Ação não concluída.')
        return redirect('url_forensicscans_list')

######## FIM ForensicScan ####################

######## NATURE OF ACTION ###################
@login_required
def natureofaction_create(request):
    template_name = 'natureofaction/form.html'
    data = {}
    data['title'] = "Cadastro dos Tipos de Ações"     
    if request.method == 'POST':        
        form = NatureOfActionForm(request.POST)
        if form.is_valid():
            form.save()                        
            return redirect('url_natureofactions_list')
        else:
            print("algo não está valido.")
    else:
        form = NatureOfActionForm()             
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def natureofactions_list(request):
    template_name = "natureofaction/list.html"
    title = "Lista dos Tipos de Ações"    
    natureofactions = NatureOfAction.objects.all()                        
    context = {
        'natureofactions': natureofactions,
        'title': title,                      
    }
    return render(request,template_name,context)

@login_required
def natureofaction_detail(request, pk=None, *args, **kwargs):
    template_name = "natureofaction/detail.html"
    natureofaction = get_object_or_404(NatureOfAction,pk=pk)
    context = {
        'natureofaction': natureofaction
    }
    return render(request, template_name, context)

@login_required
def natureofaction_edit(request, pk):    
    template_name='natureofaction/form.html'
    data = {}     
    natureofaction = get_object_or_404(NatureOfAction, pk=pk)        
    user_created = natureofaction.user_created # Esta linha faz com que o user_created não seja modificado, para mostrar quem criou esta pessoa
    form = NatureOfActionForm(request.POST or None, instance=natureofaction)
    if form.is_valid():        
        natureofaction = form.save(commit=False)
        natureofaction.user_created = user_created               
        natureofaction.save()
        return redirect('url_natureofaction_detail',pk=pk)    
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def natureofaction_delete(request,pk):
    natureofaction = get_object_or_404(NatureOfAction, pk=pk)    
    if request.method == 'GET': 
        try:               
            natureofaction.delete()
            messages.success(request, 'Ação concluída com sucesso.')
            return redirect('url_natureofactions_list')
        except IntegrityError:
            messages.warning(request, 'Não foi possível excluir. Possui dependência.')
            return redirect('url_natureofactions_list')

    else:
        messages.warning(request, 'Ação não concluída.')
        return redirect('url_natureofactions_list')


####### FIM NATURE OF ACTION ################

######## CID NUMBER ##########################

@login_required
def cidnumber_create(request):      
    id_cid_number = request.POST.get('id_cid_number',"")
    id_anamnesis_diagnosis = request.POST.get('id_anamnesis_diagnosis',"")
    id_cid_report = request.POST.get('id_cid_report',"")
    id_user_created = request.POST.get('id_user_created',"")
    id_user_updated = request.POST.get('id_user_updated',"")

    report = Report.objects.get(pk=id_cid_report)
    cid = CidNumber.objects.filter(report_id=report.id, type_cid=True).exists()
    if cid:
        type_cid = False
    else:
        type_cid = True
    results = {}
    if request.method == 'POST':
        if id_cid_number != "" and id_anamnesis_diagnosis != "" and id_cid_report != "" and id_user_created != "" and id_user_updated != "":            
            report = Report.objects.get(pk=int(id_cid_report))
            user = User.objects.get(pk=int(id_user_created))                                           
            cid = CidNumber(category=id_cid_number, description=id_anamnesis_diagnosis, report=report, type_cid = type_cid, user_created=user, user_updated=user)
            cid.save()          
            results['success'] = "Os dados foram inseridos com sucesso."
            return JsonResponse(results,status=200)

        else:
            results['error'] = "Informações incompletas. Todos os campos devem ser preenchidos."
            return JsonResponse(results,status=400)

    else:        
        results['error'] = "Requisição incorreta. Procure o administrador."
        return JsonResponse(results,status=400)

@login_required
def cidnumbers_list(request):
    template_name = "cidnumber/list.html"
    title = "Tipos de Status de Laudos"
    cidnumbers = CidNumber.objects.all()    
    context = {
        'cidnumbers': cidnumbers,
        'title': title      
    }
    return render(request,template_name,context)

@login_required
def cidnumber_detail(request, pk=None, *args, **kwargs):
    template_name = "cidnumber/detail.html"
    cidnumber = get_object_or_404(CidNumber,pk=pk)
    context = {
        'cidnumber': cidnumber
    }
    return render(request, template_name, context)

@login_required
def cidnumber_edit(request, pk):    
    template_name='cidnumber/form.html'
    data = {} 
    cidnumber = get_object_or_404(CidNumber, pk=pk)        
    user_created = cidnumber.user_created # Esta linha faz com que o user_created não seja modificado, para mostrar quem criou esta pessoa
    form = CidNumberForm(request.POST or None, instance=cidnumber)
    if form.is_valid():        
        cidnumber = form.save(commit=False)
        cidnumber.user_created = user_created               
        cidnumber.save()
        return redirect('url_cidnumber_detail',pk=pk)    
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def cidnumber_delete(request,pk):
    request.session['tabs'] = 'diagnostico-tab'
    cidnumber = get_object_or_404(CidNumber, pk=pk)  
    report_id = cidnumber.report_id
    if request.method == 'GET':
        if not cidnumber.type_cid:                 
            cidnumber.delete()            
            messages.success(request, 'Ação concluída com sucesso.')            
        else:
            count = CidNumber.objects.filter(report_id=report_id,type_cid=False).count()
            if count > 0:
                messages.warning(request, 'Para excluir o CID Primário deverá deletar os secundários.')
            else:
                cidnumber.delete()            
                messages.success(request, 'Ação concluída com sucesso.')                        
    else:                
        messages.warning(request, 'Ação não concluída.')
    
    return redirect('url_report_edit', report_id)
######## FIM CID NUMBER ######################


####### REPORT STATUS #######################

@login_required
def reportstatus_create(request):
    template_name = 'reportstatus/form.html'
    data = {}
    data['title'] = "Cadastro Status do Laudo"        
    if request.method == 'POST':        
        form = ReportStatusForm(request.POST)
        if form.is_valid():
            form.save()                        
            return redirect('url_reportstatus_list')
        else:
            print("algo não está valido.")
    else:
        form = ReportStatusForm()             
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def reportstatus_list(request):
    template_name = "reportstatus/list.html"
    title = "Tipos de Status de Laudos"
    reportstatus = ReportStatus.objects.all()    
    context = {
        'reportstatus': reportstatus,
        'title': title      
    }
    return render(request,template_name,context)

@login_required
def reportstatus_detail(request, pk=None, *args, **kwargs):
    template_name = "reportstatus/detail.html"
    reportstatus = get_object_or_404(ReportStatus,pk=pk)
    context = {
        'reportstatus': reportstatus
    }
    return render(request, template_name, context)

@login_required
def reportstatus_edit(request, pk):    
    template_name='reportstatus/form.html'
    data = {} 
    reportstatus = get_object_or_404(ReportStatus, pk=pk)        
    user_created = reportstatus.user_created # Esta linha faz com que o user_created não seja modificado, para mostrar quem criou esta pessoa
    form = ReportStatusForm(request.POST or None, instance=reportstatus)
    if form.is_valid():        
        reportstatus = form.save(commit=False)
        reportstatus.user_created = user_created               
        reportstatus.save()
        return redirect('url_reportstatus_detail',pk=pk)    
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def reportstatus_delete(request,pk):
    reportstatus = get_object_or_404(ReportStatus, pk=pk)    
    if request.method == 'GET':        
        reportstatus.delete()
        messages.success(request, 'Ação concluída com sucesso.')
        return redirect('url_reportstatus_list')
    else:
        messages.warning(request, 'Ação não concluída.')
        return redirect('url_reportstatus_list')

###### FIM REPORT STATUS ###################

###### DISCUSSION CONCLUSION ##############

@login_required
def discussionconclusion_create(request):
    template_name = 'discussionconclusion/form.html'
    data = {}
    data['title'] = "Cadastro dos Modelos de Discussão e Conclusão"        
    if request.method == 'POST':        
        form = DiscussionConclusionForm(request.POST,department_id=get_department(request.user.username))
        if form.is_valid():
            form.save()                        
            return redirect('url_discussionconclusions_list')
        else:
            print("algo não está valido.")
    else:
        form = DiscussionConclusionForm(department_id=get_department(request.user.username))             
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def discussionconclusions_list(request): 
    template_name = "discussionconclusion/list.html"
    title = "Modelos de Discussão e Conclusão"
    #discussionconclusions = DiscussionConclusion.objects.all() 
    discussionconclusions = DiscussionConclusion.objects.filter(doctor__profile_person_type__department_id__in=get_department(request.user.username))               
    context = {
        'discussionconclusions': discussionconclusions,
        'title': title      
    }
    return render(request,template_name,context)

@login_required
def discussionconclusion_detail(request, pk=None, *args, **kwargs):
    template_name = "discussionconclusion/detail.html"
    discussionconclusion = get_object_or_404(DiscussionConclusion,pk=pk)
    context = {
        'discussionconclusion': discussionconclusion
    }
    return render(request, template_name, context)

@login_required
def discussionconclusion_edit(request, pk):    
    template_name='discussionconclusion/form.html'
    title = "Editando campos do Modelo - Discussão e Conclusão"
    data = {} 
    data['title'] = title
    discussionconclusion = get_object_or_404(DiscussionConclusion, pk=pk)        
    user_created = discussionconclusion.user_created # Esta linha faz com que o user_created não seja modificado, para mostrar quem criou esta pessoa
    form = DiscussionConclusionForm(request.POST or None, instance=discussionconclusion, department_id=get_department(request.user.username))
    if form.is_valid():        
        discussionconclusion = form.save(commit=False)
        discussionconclusion.user_created = user_created               
        discussionconclusion.save()
        return redirect('url_discussionconclusion_detail',pk=pk)    
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def discussionconclusion_delete(request,pk):
    discussionconclusion = get_object_or_404(DiscussionConclusion, pk=pk)    
    if request.method == 'GET':        
        discussionconclusion.delete()
        messages.success(request, 'Ação concluída com sucesso.')
        return redirect('url_discussionconclusions_list')
    else:
        messages.warning(request, 'Ação não concluída.')
        return redirect('url_discussionconclusions_list')

@login_required
def discussionconclusion_delete_all(request):
    context = []
    context = request.GET.get("ff","")
    context = context.replace("[","").replace("]","").replace('\"','')
    context = context.split(",")
    context = [int(x) for x in context]      
    if context:                
        b = DiscussionConclusion.objects.filter(id__in=context)
        b.delete()
        results = []         
        json = {}
        json['message'] = "Deletado"        
        results.append(json)
        return JsonResponse(results, safe=False)        
    else:
        json['message'] = "Não deletado"        
        results.append(json)
        return JsonResponse(results, safe=False)

###### FIM DISCUSSION CONCLUSION ##########

###### REPORT #############################

@login_required
def report_create(request):
    template_name = 'report/form.html'
    data = {}
    #Isso é um marcador das TABS para saber que estou criando
    data['tabs'] = "preambulo-tab"
    data['create'] = "create"
    data['title'] = "Criação do Laudo" 
    type_items = TypeItem.objects.all()
    data['type_items'] = type_items          
    if request.method == 'POST':        
        form = ReportForm(request.POST,department_id=get_department(request.user.username),user_id=request.user.id)
        if form.is_valid():
            form.save()                        
            return redirect('url_reports_list')
        else:
            print("Form errors",form.errors)
            print("algo não está valido.")
    else:
        form = ReportForm(department_id=get_department(request.user.username),user_id=request.user.id)             
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def reports_list(request):
    status = request.GET.get("status",0)
    print("Valor status",status)
    template_name = "report/list.html"
    title = "Lista do Laudos"
    if not status:
        reports = Report.objects.filter(doctor__profile_person_type__department_id__in=get_department(request.user.username))    
    else:
        reports = Report.objects.filter(report_status=int(status),doctor__profile_person_type__department_id__in=get_department(request.user.username))            
    context = {
        'reports': reports,
        'title': title      
    }
    return render(request,template_name,context)

@login_required
def report_detail(request, pk=None, *args, **kwargs):
    template_name = "report/detail.html"
    report = get_object_or_404(Report,pk=pk)
    context = {
        'report': report
    }
    return render(request, template_name, context)

# Função usada apenas para retornar a TAB e identificar qual aba o usuário está editando. A função é usada no EDIT
def get_report_tab(tab):    
    tabs = {
        "preambulo-tab": "preambulo-tab",
        "perito-tab": "perito-tab",
        "anamnese-tab": "anamnese-tab",
        "documentos-tab": "documentos-tab",
        "diagnostico-tab": "diagnostico-tab",
        "discussao-tab": "discussao-tab",
        "conclusao-tab": "conclusao-tab",
        "quesito-tab": "quesito-tab",
        "imp-tab": "imp-tab",
    }
    return tabs.get(tab,"1")   

@login_required
def report_edit(request, pk):   
    #Isso é um marcador das TABS. Quando o perito salvar, vai redirecionar para a aba/tab que ele estava usando    
    if request.method == "POST":        
        tabs = get_report_tab(request.POST["value_tab"])
        if not tabs.isdigit():
            request.session['tabs'] = tabs                           
    
    template_name='report/form.html'
    title = "Criação do Laudo" 
    report = get_object_or_404(Report, pk=pk)
    type_items = TypeItem.objects.all()# Isso é para o Mostrar os quesitos possíveis    
    medicaldocuments = MedicalDocument.objects.filter(report_id=report.id).order_by('document')# Isso é para o Mostrar os documentos possíveis    
    cidnumbers = CidNumber.objects.filter(report_id=report.id).exists()# Isso é para o Mostrar os cid deste laudo    
    if cidnumbers:# Se entrar é pq existe diagnostico para este laudo
        cidnumbers = CidNumber.objects.filter(report_id=report.id)
        cid_primario = CidNumber.objects.get(report_id=report.id, type_cid=True)
        #Verifica se existem modelo para este diagnóstico para este Perito
        discon = DiscussionConclusion.objects.filter(cid_number=cid_primario.category, doctor=report.doctor).exists()
        if discon:
            discon = DiscussionConclusion.objects.get(cid_number=cid_primario.category, doctor=report.doctor)
        else:
            discon = "Não existe modelo para este diagnóstico."    
    else:
        cidnumbers = {}
        discon = "Não existe diagnóstico para este laudo."
    data = {
        'title': title,
        'edit': 'edit',# Marcador para saber que estou editando
        'tabs': request.session.get('tabs',''),  
        'report': report,  
        'type_items': type_items,
        'medicaldocuments': medicaldocuments,
        'cidnumbers': cidnumbers,
        'discon': discon,
    } 
    # As duas linhas abaixo servem para gerar o relatório que será inserido no ckeditor na tab Visualização
    #LEMBRAR: Ele vai mostrar o conteudo sempre atrasado, pois está carregando os valores antes de atualizar, com isso fica com os valores antigos
    #por isso que fiz um submit toda vez que clicar na tab Visualização
    context_imp = print_ckeditor(report.id)
    valores = render_to_string('report/print_report_ckeditor.html',context_imp)        
    
    user_created = report.user_created # Esta linha faz com que o user_created não seja modificado, para mostrar quem criou esta pessoa
    form = ReportForm(request.POST or None, instance=report, department_id=get_department(request.user.username), user_id=request.user.id)
    if form.is_valid():        
        report = form.save(commit=False)
        report.user_created = user_created# O usuario que criação não deve ser atualizado
        report.impress = valores# Essa campo é somente leitura no form, então ele é inserido aqui com os valores do Laudo               
        report.save()
        return redirect('url_report_edit',pk=pk)   
    
    data['tabs'] =  request.session.get('tabs','')
    data['form'] = form   
    return render(request,template_name,data)

@login_required
def report_delete(request,pk):
    report = get_object_or_404(Report, pk=pk)    
    if request.method == 'GET':        
        report.delete()
        messages.success(request, 'Ação concluída com sucesso.')
        return redirect('url_reports_list')
    else:
        messages.warning(request, 'Ação não concluída.')
        return redirect('url_reports_list')

# Esta função serve para pegar a Circuntância e o Objetivo e colocar dentro dos cards do template, isso VIA AJAX
@login_required
def get_report_location_objective(request):
    print("Valor da request get", request.GET)
    ff = request.GET.get('ff','')
    context = {}       
    marc = 0    
    if ff:         
        location_objective = LocationObjective.objects.get(pk=int(ff)) 
        print("Valor do location_objective",location_objective)
        marc = 1
        context['location_objective'] = location_objective 
        context['doctor'] = Doctor.objects.get(id=location_objective.doctor_id)
    
    context['marc'] = marc
    return render(request, 'report/_location_objective.html',context)
    
# AJAX para copiar ANAMNESE, DISCUSSÃO E CONCLUSÃO
@login_required
def forensic_copy_report(request):    
    ff = request.GET.get('ff','')
    id_tab = request.GET.get('tab','')
    doctor = request.GET.get('perito','')   
    id_anamnese = request.GET.get('anamnese','')   
    context = {}       
    marc = 0    
    if ff:                          
        if id_tab == "anamnese-tab":            
            forensicscan = ForensicScan.objects.get(pk=int(ff))
            if id_anamnese == "id_anamnesis_family":
                context['forensicscan'] = forensicscan.anamnesis_family
            elif id_anamnese == "id_anamnesis_professional":
                context['forensicscan'] = forensicscan.anamnesis_professional
            elif id_anamnese == "id_anamnesis_conditions":
                context['forensicscan'] = forensicscan.anamnesis_conditions
            elif id_anamnese == "id_anamnesis_history":
                context['forensicscan'] = forensicscan.anamnesis_history
            elif id_anamnese == "id_anamnesis_personal_background":
                context['forensicscan'] = forensicscan.anamnesis_personal_background
            elif id_anamnese == "id_anamnesis_family_background":
                context['forensicscan'] = forensicscan.anamnesis_family_background
            elif id_anamnese == "id_anamnesis_general_exam":
                context['forensicscan'] = forensicscan.anamnesis_general_exam
            elif id_anamnese == "id_anamnesis_mental_exam":
                context['forensicscan'] = forensicscan.anamnesis_mental_exam

        elif id_tab == "discussao-tab":            
            copy = DiscussionConclusion.objects.filter(cid_number=ff, doctor=doctor).exists()
            if copy:
                copy = DiscussionConclusion.objects.get(cid_number=ff, doctor=doctor)
                context['forensicscan'] = copy.discussion                
            else:
                print("caiu aqui no else do copy")
                context['forensicscan'] = "Não existe modelo para este diagnóstico." 
            
        elif id_tab == "conclusao-tab":
            copy = DiscussionConclusion.objects.filter(cid_number=ff, doctor=doctor).exists()
            if copy:
                copy = DiscussionConclusion.objects.get(cid_number=ff, doctor=doctor)
                context['forensicscan'] = copy.conclusion          
            else:
                context['forensicscan'] = 'Não existe modelo para este diagnóstico.' 
            

        else:
            context['forensicscan'] = "Erro....."                    
    
    return render(request, 'report/_forensic_copy_report.html',context)

#Função que valida se o Laudo pode ser aprovado
@login_required
def valid_report_item(request,pk):    
    
    report = Report.objects.get(pk=pk)    
    valid = []
    if report.authority_requesting == "":
        valid.append("Autoridade Requisitante")
    if report.nature_of_action == "":
        valid.append("Natureza da Ação")
    if report.process_number == "":
        valid.append("Número do Processo")
    if report.autor == "":
        valid.append("Autor")
    if report.proficient == "":
        valid.append("Periciando")
    if report.date_report == "":
        valid.append("Data da Perícia")
    if report.location_objective == None:
        valid.append("Circunstacias e Objetivos")
    if report.forensic_scan == None:
        valid.append("Circunstância do Exame")
    if report.anamnesis_history == "":
        valid.append("Anamnese: História Pregressa da Doença Atual")
    if report.anamnesis_personal_background == "":
        valid.append("Anamnese: Antecedentes Patológicos Pessoais")
    if report.anamnesis_family_background == "":
        valid.append("Anamnese: Antecedentes Patológicos Familiares")
    if report.anamnesis_general_exam == "":
        valid.append("Anamnese: Exame Físico Geral")
    if report.anamnesis_mental_exam == "":
        valid.append("Anamnese: Exame do Estado Mental")        
    if report.discussion == "":
        valid.append("Discussão")
    if report.conclusion == "":
        valid.append("Conclusão")

    #Verifica se tem Diagnóstico e tem que ser primario
    cid_number = CidNumber.objects.filter(report_id=report.id, type_cid=True)
    
    aux2 = "Os seguintes quesitos não estão preenchidos: "
    marc = Item2.objects.filter(report_id=report.id).exists()
    marc2 = -1    
    lista_type_item = []
    #Marcados para saber se existem quesitos para este laduo
    if marc:        
        items = Item2.objects.filter(report_id=report.id)# Pega todos os quesitos de laudo
        for item in items:
            lista_type_item.append(item.type_item_id)

        type_items = TypeItem.objects.exclude(id__in=lista_type_item)#Verifica qual tipo de quesito não existe neste laudo      
        for type_item in type_items:                                          
            if type_item.id == 1:
                aux2 += type_item.name + " "
                marc2 = 1
            elif type_item.id == 2:
                aux2 += type_item.name + " "
                marc2 = 1

    # Testa se é o perito cadastrado que está aprovando o laudo
    if report.doctor.profile_person_type.person.cpf == request.user.username:
        if len(valid) > 0:         
            aux = "Os campos não validados são: "
            for n in valid:
                aux += n + " "
            if marc2 == 1:
                aux += aux2                                         
            messages.warning(request, 'Existem campos não preenchidos.Não é possível validar. {}'.format(aux))                     
        
         #INFORMA SE TEM DIAGNÓSTICO
        elif not cid_number:            
            messages.warning(request, 'Aprovação negada. Diagnóstico não cadastrado.')

        #INFORMA SE EXISTEM QUESITOS NESTE LAUDO
        elif not marc:
            messages.warning(request, 'Aprovação negada. Não existem quesitos para este Laudo.')
        #NÃO APAGAR ISSO. ISSO VALIDAR E INFORMA QUAL QUESITO FALTA PREENCHER. OS QUESITOS OBRIGATÓRIOS SÃO JUIZ E PROCURADORIA. OS OUTROS SÃO OPCIONAIS
        elif marc2 == 1:
            messages.warning(request, 'Existem campos não preenchidos.Não é possível validar. {}'.format(aux2))

        else:        
            aprovar = ReportStatus.objects.get(id=2)# Esse é o id para aprovar
            report.report_status = aprovar # Deixa o status validado
            report.save()
            messages.success(request, 'Ação concluída com sucesso.')
    else:
        messages.warning(request, 'Aprovação negada. Somente o perito do Laudo pode aprovar.')
    
    return redirect('url_reports_list')

#Função que cancela o laudo
@login_required
def cancelar_report(request,pk):
    report = Report.objects.get(pk=pk)
    cancelar = ReportStatus.objects.get(id=3)# Esse é o id para cancelar
    report.report_status = cancelar # Deixa o status cancelado
    report.save()
    messages.success(request, 'Ação concluída com sucesso.')
    
    return redirect('url_reports_list')

@login_required
def print_docx(request,pk):
    context = print_report(pk)
    #Faz parte do modulo docxtpl    
    doc = DocxTemplate("files/teste_2.docx")          
   

    context['report'].date_report = context['report'].date_report.strftime('%d-%m-%Y')
    context['report'].proficient.dt_birthday = context['report'].proficient.dt_birthday.strftime('%d-%m-%Y')
    
    #Calculo da Idade
    agora = datetime.now()
    niver = context['report'].proficient.dt_birthday.split("-")    
    aniversario = datetime(int(niver[2]),int(niver[1]),int(niver[0]))
    idade=(agora - aniversario)
    idade = (idade.days)/365
    context['idade'] = int(idade)

    #Utiliza a biblioteca BeautifulSoup para retirar as tags html e inserir no DOCX. O PDF não tem esse problema
    soup = BeautifulSoup(context['report'].anamnesis_history,'lxml')
    context['report'].anamnesis_history = soup.get_text().strip()

    soup = BeautifulSoup(context['report'].anamnesis_personal_background,'lxml')
    context['report'].anamnesis_personal_background = soup.get_text().strip()

    soup = BeautifulSoup(context['report'].anamnesis_family_background,'lxml')
    context['report'].anamnesis_family_background = soup.get_text().strip()

    soup = BeautifulSoup(context['report'].anamnesis_general_exam,'lxml')
    context['report'].anamnesis_general_exam = soup.get_text().strip()

    soup = BeautifulSoup(context['report'].anamnesis_mental_exam,'lxml')
    context['report'].anamnesis_mental_exam = soup.get_text().strip()

    soup = BeautifulSoup(context['report'].discussion,'lxml')
    context['report'].discussion = soup.get_text().strip()

    soup = BeautifulSoup(context['report'].conclusion,'lxml')
    context['report'].conclusion = soup.get_text().strip()

    soup = BeautifulSoup(context['item_juiz'],'lxml')
    context['item_juiz'] = soup.get_text().strip()

    soup = BeautifulSoup(context['item_procuradoria'],'lxml')
    context['item_procuradoria'] = soup.get_text().strip()

    soup = BeautifulSoup(context['item_advogado'],'lxml')
    context['item_advogado'] = soup.get_text().strip()

    soup = BeautifulSoup(context['item_periciando'],'lxml')
    context['item_periciando'] = soup.get_text().strip()

    context['report'].location_objective.forensic_scan = context['report'].location_objective.forensic_scan.strip() 
    context['report'].location_objective.goal = context['report'].location_objective.goal.strip()
    context['doctor'].course = context['doctor'].course.strip()
  
    # documento = doc.new_subdoc()
    # if context['medicaldocuments']:        
    #     for medicaldocument in context['medicaldocuments']:            
    #         documento.add_paragraph(medicaldocument.document.strip("\n"))                        
    #         print("aqui")          
    #     context['documents'] = documento
    # else:
    #     context['documents'] = "Não há."
        

    documento = ""
    if context['medicaldocuments']:
        for medicaldocument in context['medicaldocuments']:
            print("Documento",medicaldocument.document)
            documento += medicaldocument.document + "\n"           
        context['documents'] = documento
    else:
        context['documents'] = "Não há."

    doc.render(context, autoescape=True)
    doc.save("files/teste1.docx")
    #Começa a montagem para download    
    filename = "Laudo{}.{}".format(datetime.now().strftime('%d_%m_%Y'),'docx') # colocar aqui a data    
    arq = open('files/teste1.docx','rb').read()    
    content = "attachment; filename={}".format(filename)
    #response = HttpResponse(arq,content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response = HttpResponse(arq,content_type='application/vnd.ms-word')
    response['Content-Disposition'] = content        
    return response

@login_required
def print_pdf(request,pk):
    template_name = 'report/report_view_after_save.html'
    title = 'Visualização do Laudo - Aprovado'
    report = Report.objects.get(pk=pk)    
    valores = report.impress
    context = {
        'view': valores,
        'title': title,
    }
        #Essa função fica em um arquivo chamado utils.py nesta App Report
    pdf = render_to_pdf('report/print_report_pdf.html',context)  
    
    if pdf:
        response = HttpResponse(pdf, content_type='application/pdf')
        filename = "Laudo{}.{}".format(datetime.now().strftime('%d_%m_%Y'),'pdf') # colocar aqui a data
        content = "inline; filename={}".format(filename)
        download = request.GET.get("download","")
        if download:
            content = "attachment; filename={}".format(filename)
        response['Content-Disposition'] = content
        return response
    return HttpResponse("Não encontrado")
    #response['Content-Disposition'] = 'attachment; filename="somefilename.pdf"'

@login_required
def report_view_after_save(request,pk):
    template_name = 'report/report_view_after_save.html'
    title = 'Visualização do Laudo - Aprovado'
    report = Report.objects.get(pk=pk)    
    valores = report.impress
    context = {
        'view': valores,
        'title': title,
    }
    return render(request,template_name,context)

def print_ckeditor(pk):
    context = print_report(pk)

    #Calculo da Idade
    agora = datetime.now()
    niver = context['report'].proficient.dt_birthday.strftime('%d-%m-%Y')
    niver = niver.split("-")   
    aniversario = datetime(int(niver[2]),int(niver[1]),int(niver[0]))
    idade=(agora - aniversario)
    idade = (idade.days)/365
    context['idade'] = int(idade)    
    
    if context['report'].location_objective:# Senão existir ele gera um erro
        soup = BeautifulSoup(context['report'].location_objective.forensic_scan,'lxml')
        context['report'].location_objective.forensic_scan = soup.get_text().strip()
        soup = BeautifulSoup(context['report'].location_objective.goal,'lxml')
        context['report'].location_objective.goal = soup.get_text().strip()
    
    
    documento = ""
    if context['medicaldocuments']:
        for medicaldocument in context['medicaldocuments']:
            documento += "{} <br/>".format(medicaldocument.document)           
        context['documents'] = documento
    else:
        context['documents'] = "Não há.<br/>"

    cid = ""
    if context['cid_numbers']:
        for cid_number in context['cid_numbers']:
            cid += "-{} <br/>".format(cid_number.description)           
        context['cid_numbers'] = cid
    else:
        context['cid_numbers'] = "Não há.<br/>"

    return context

#Função que pega os valores para inserir no template para impressão
def print_report(pk):
    report = Report.objects.get(pk=pk)
    doctor = Doctor.objects.get(pk=report.doctor.id)    
    medicaldocuments = MedicalDocument.objects.filter(report_id=report.id).order_by('document')
    marc = Item2.objects.filter(report_id=report.id).exists()
    cid_numbers = CidNumber.objects.filter(report_id=report.id)
     
    item_juiz = ""
    item_procuradoria  = ""
    item_advogado = ""    
    item_periciando = ""

    if marc:        
        items = Item2.objects.filter(report_id=report.id).order_by("id")
        for item in items:
            #Na tabela type_item tem que ser assim: id: 1 - Quesitos do juiz, 2 - Procuradoria, 3 - ADvogado e 4 - Periciando. Senão for assim, a lógica abaixo vai dar erro.
            if item.type_item_id == 1:               
                item_juiz += "{} \n".format(item.question)
                item_juiz += "{} \n".format(item.answer)            
                item_juiz += "\n"
            elif item.type_item_id == 2:
                item_procuradoria += "{} \n".format(item.question)
                item_procuradoria += "{} \n".format(item.answer)
                item_procuradoria += "\n"
            elif item.type_item_id == 3:
                item_advogado += "{} \n".format(item.question)
                item_advogado += "{} \n".format(item.answer) 
                item_advogado += "\n"                 
            elif item.type_item_id == 4:
                item_periciando += "{} \n".format(item.question)
                item_periciando += "{} \n".format(item.answer) 
                item_periciando += "\n"                  
        
    context = { 
        'report' : report,
        'item_juiz': item_juiz,
        'item_procuradoria': item_procuradoria,
        'item_advogado': item_advogado,
        'item_periciando': item_periciando,
        'doctor': doctor,
        'medicaldocuments': medicaldocuments,
        'cid_numbers': cid_numbers,    
    }
    
    #Utiliza a biblioteca BeautifulSoup para retirar as tags html e inserir no DOCX. O PDF não tem esse problema
    soup = BeautifulSoup(context['report'].anamnesis_family,'lxml')
    context['report'].anamnesis_family = soup.get_text().strip()
    soup = BeautifulSoup(context['report'].anamnesis_professional,'lxml')
    context['report'].anamnesis_professional = soup.get_text().strip()
    soup = BeautifulSoup(context['report'].anamnesis_conditions,'lxml')
    context['report'].anamnesis_conditions = soup.get_text().strip()    
    soup = BeautifulSoup(context['report'].anamnesis_history,'lxml')
    context['report'].anamnesis_history = soup.get_text().strip()
    soup = BeautifulSoup(context['report'].anamnesis_personal_background,'lxml')
    context['report'].anamnesis_personal_background = soup.get_text().strip()
    soup = BeautifulSoup(context['report'].anamnesis_family_background,'lxml')
    context['report'].anamnesis_family_background = soup.get_text().strip()
    soup = BeautifulSoup(context['report'].anamnesis_general_exam,'lxml')
    context['report'].anamnesis_general_exam = soup.get_text().strip()
    soup = BeautifulSoup(context['report'].anamnesis_mental_exam,'lxml')
    context['report'].anamnesis_mental_exam = soup.get_text().strip()
    
    return context

#Lista os Laudo aprovados e cancelados
@login_required
def search_reports(request,status):
    template_name = 'report/searchs_report.html'
    title = "Laudos aprovados e cancelados"
    reports = Report.objects.filter(report_status=status)
    context = {
        'reports': reports,
        'title': title,
    } 
    return render(request,template_name,context)


##### FIM DO REPORT #######################

##### MEDICAL DOCUMENTS ###################

@login_required
def medicaldocument_create(request):
    ordem = {
        '0': 'a)', '1': 'b)', '2': 'c)', '3': 'd)', '4': 'e)', '5': 'f)', '6': 'g)', '7': 'h)', '8': 'i)', '9': 'j)',
        '10': 'k)', '11': 'l)', '12': 'm)', '13': 'n)', '14': 'p)', '15': 'q)', '16': 'r)', '17': 's)', '18': 't)', '19': 'u)',
        '20': 'u)', '21': 'v)', '22': 'w)','23': 'x)', '24': 'y)', '25': 'z)',
    }
    if request.method == 'POST':               
        report_id = request.POST.get('id_document_report',"")
        tipo = request.POST.get('id_tipo',"")
        dt = request.POST.get('id_data',"")
        medico = request.POST.get('id_medico',"")
        sexo = request.POST.get('id_medico_sexo',"")
        cid = request.POST.get('id_cid',"")
        user_created = request.POST.get('id_user_created',"")
        user_updated = request.POST.get('id_user_updated',"")
        user_updated = request.POST.get('id_user_updated',"")

        textarea_document = request.POST.get('id_textarea_document',"")
        #Faço o split para tornar cada CID uma lista para fazer a consulta logo abaixo
        aux = cid.split(",")
        
        results = {}
        if int(tipo) == 1:
            if tipo != "" and dt != "" and medico != "" and cid != "" and user_created != "" and user_updated != "" and report_id != "":            
                report = Report.objects.get(pk=int(report_id))
                user = User.objects.get(pk=int(user_created))                                               
                cid_completo = Cid10.objects.filter(category__in=aux)  
                cids = ""
                cont = cid_completo.count() 
                cont2 = cont               
                for i in cid_completo:
                    if cont < 2:
                        cids += "CID 10 {} - {}".format(i.category,i.description)
                    else:
                        cids += "CID 10 {} - {}, ".format(i.category,i.description)
                    cont -= 1
                count = MedicalDocument.objects.filter(report = report).count()
                print("Valor do count",count)
                if sexo == 'M' and int(cont2) == 1:                                         
                    document = "{} Atestado emitido em {} pelo médico psiquiatra {},  com o seguinte diagnóstico: {}.".format(ordem.get(str(count),""), dt.strip(),medico.strip(), cids)                
                elif sexo == 'M' and int(cont2) > 1:                                         
                    document = "{} Atestado emitido em {} pelo médico psiquiatra {},  com os seguintes diagnósticos: {}.".format(ordem.get(str(count),""), dt.strip(),medico.strip(), cids)                
                elif sexo == 'F' and int(cont2) == 1:                                         
                    document = "{} Atestado emitido em {} pela médica psiquiatra {},  com o seguinte diagnóstico: {}.".format(ordem.get(str(count),""), dt.strip(),medico.strip(), cids)                
                else:
                    document = "{} Atestado emitido em {} pela médica psiquiatra {},  com os seguintes diagnósticos: {}.".format(ordem.get(str(count),""), dt.strip(),medico.strip(), cids)                
                md = MedicalDocument(report=report, document=document, user_created=user, user_updated=user)
                md.save()          
                results['success'] = "Os dados foram inseridos com sucesso."
                return JsonResponse(results,status=200)

            else:
                results['error'] = "Informações incompletas. Todos os campos devem ser preenchidos."
                return JsonResponse(results,status=400)
        else:
            if textarea_document != "" and user_created != "" and user_updated != "" and report_id != "":
                report = Report.objects.get(pk=int(report_id))
                user = User.objects.get(pk=int(user_created))
                document = textarea_document.strip()                
                md = MedicalDocument(report=report, document=document, user_created=user, user_updated=user)
                md.save()
                results['success'] = "Os dados foram inseridos com sucesso."
                return JsonResponse(results,status=200)
            else:
                results['error'] = "Informações incompletas. Todos os campos devem ser preenchidos."
                return JsonResponse(results,status=400)

            
    else:        
        results['error'] = "Requisição incorreta. Procure o administrador."
        return JsonResponse(results,status=400)

@login_required
def medicaldocuments_list(request):
    template_name = "report/custom_edit_documentos_medicos.html"
    title = "Lista de Documentos Médicos do Periciando(a)"        
    medicaldocuments = MedicalDocument.objects.all()     
    #typeitems = TypeItem.objects.filter(user_created = request.user.id)  
    context = {
        'medicaldocuments': medicaldocuments,
        'title': title,                      
    }
    return render(request,template_name,context)

@login_required
def medicaldocument_delete(request,pk):
    medicaldocument = get_object_or_404(MedicalDocument, pk=pk)    
    report_id = medicaldocument.report_id
    if request.method == 'GET':        
        medicaldocument.delete()        
        request.session['tabs'] = 'documentos-tab'
        return redirect('url_report_edit', report_id)
    else:        
        request.session['tabs'] = 'documentos-tab'
        return redirect('url_report_edit', report_id)
##### FIM MEDICAL DOCUMENTS ###############

###### TYPE ITEM (TIPOS DE QUESITOS) #####

@login_required
def typeitem_create(request):
    template_name = 'typeitem/form.html'
    data = {}
    data['title'] = "Cadastro das Autoridades que solicitam os Quesitos/Perguntas"
    if request.method == 'POST':        
        form = TypeItemForm(request.POST)
        if form.is_valid():
            form.save()                        
            return redirect('url_typeitems_list')
        else:
            print("algo não está valido.")
    else:
        form = TypeItemForm()             
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def typeitems_list(request):
    template_name = "typeitem/list.html"
    title = "Lista das Autoridades que solicitam Quesitos"        
    typeitems = TypeItem.objects.all()     
    #typeitems = TypeItem.objects.filter(user_created = request.user.id)  
    context = {
        'typeitems': typeitems,
        'title': title,                      
    }
    return render(request,template_name,context)

@login_required
def typeitem_detail(request, pk=None, *args, **kwargs):
    template_name = "typeitem/detail.html"
    typeitem = get_object_or_404(TypeItem,pk=pk)
    context = {
        'typeitem': typeitem
    }
    return render(request, template_name, context)

@login_required
def typeitem_edit(request, pk):    
    template_name='typeitem/form.html'
    data = {}    
    typeitem = get_object_or_404(TypeItem, pk=pk)    
    user_created = typeitem.user_created # Esta linha faz com que o user_created não seja modificado, para mostrar quem criou esta pessoa
    form = TypeItemForm(request.POST or None, instance=typeitem)
    if form.is_valid():        
        typeitem = form.save(commit=False)
        typeitem.user_created = user_created               
        typeitem.save()
        return redirect('url_typeitem_detail',pk=pk)    
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def typeitem_delete(request,pk):
    typeitem = get_object_or_404(TypeItem, pk=pk)    
    if request.method == 'GET':        
        typeitem.delete()
        messages.success(request, 'Ação concluída com sucesso.')
        return redirect('url_typeitems_list')
    else:
        messages.warning(request, 'Ação não concluída.')
        return redirect('url_typeitems_list')

####### FIM TYPE ITEM ###################

##### TypeItemByNatureOfAction(TIPOS DE QUESITOS POR NATUREZA DE AÇÃO) ########

@login_required
def typeitembynatureofaction_create(request):
    print("Valores do request",request.POST)
    template_name = 'typeitembynatureofaction/form.html'
    data = {}
    data['title'] = "Cadastro dos Tipos de Quesitos"             
    if request.method == 'POST':        
        form = TypeItemByNatureOfActionForm(request.POST, department_id=get_department(request.user.username))
        if form.is_valid():
            form.save()                        
            return redirect('url_typeitembynatureofactions_list')
        else:
            print("algo não está valido.",form.errors)
    else:
        form = TypeItemByNatureOfActionForm(department_id=get_department(request.user.username))             
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def typeitembynatureofactions_list(request):
    template_name = "typeitembynatureofaction/list.html"
    title = "Lista dos Modelos de Quesitos"    
    #typeitembynatureofactions = TypeItemByNatureOfAction.objects.all()                            
    typeitembynatureofactions = TypeItemByNatureOfAction.objects.filter(doctor__profile_person_type__department__in = get_department(request.user.username))       
    context = {
        'typeitembynatureofactions': typeitembynatureofactions,
        'title': title,                            
    }
    return render(request,template_name,context)

@login_required
def typeitembynatureofaction_detail(request, pk=None, *args, **kwargs):
    template_name = "typeitembynatureofaction/detail.html"
    typeitembynatureofaction = get_object_or_404(TypeItemByNatureOfAction,pk=pk)
    context = {
        'typeitembynatureofaction': typeitembynatureofaction
    }
    return render(request, template_name, context)

@login_required
def typeitembynatureofaction_edit(request, pk):    
    template_name='typeitembynatureofaction/form.html'
    data = {}      
    typeitembynatureofaction = get_object_or_404(TypeItemByNatureOfAction, pk=pk)        
    user_created = typeitembynatureofaction.user_created # Esta linha faz com que o user_created não seja modificado, para mostrar quem criou esta pessoa
    form = TypeItemByNatureOfActionForm(request.POST or None, instance=typeitembynatureofaction,department_id=get_department(request.user.username))
    if form.is_valid():        
        typeitembynatureofaction = form.save(commit=False)
        typeitembynatureofaction.user_created = user_created               
        typeitembynatureofaction.save()
        return redirect('url_typeitembynatureofaction_detail',pk=pk)    
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def typeitembynatureofaction_delete(request,pk):
    typeitembynatureofaction = get_object_or_404(TypeItemByNatureOfAction, pk=pk)    
    if request.method == 'GET':        
        typeitembynatureofaction.delete()
        messages.success(request, 'Ação concluída com sucesso.')
        return redirect('url_typeitembynatureofactions_list')
    else:
        messages.warning(request, 'Ação não concluída.')
        return redirect('url_typeitembynatureofactions_list')

#Função que deleta varios quesito de uma vez,através do checkbox e ajax
@login_required
def typeitembynatureofaction_delete_all(request):
    context = []
    context = request.GET.get("ff","")
    context = context.replace("[","").replace("]","").replace('\"','')
    context = context.split(",")
    context = [int(x) for x in context]
    print("Context transformado", context)
    if context:
        print("Valor do context",context)
        b = TypeItemByNatureOfAction.objects.filter(id__in=context)
        b.delete()
        results = []         
        json = {}
        json['message'] = "Deletado"        
        results.append(json)
        return JsonResponse(results, safe=False)        
    else:
        json['message'] = "Não deletado"        
        results.append(json)
        return JsonResponse(results, safe=False)

##### FIM TTypeItemByNatureOfAction #######################

###### ITEM 2 ################################
@login_required
def itembynatureaction_copy_item(request):
    ff = request.GET.get('ff','')
    id_tab = request.GET.get('id','')   
    context = {}       
    marc = 0    
    if ff:         
        question = TypeItemByNatureOfAction.objects.get(pk=int(ff))         
        context['question'] = question.question                   
    else:
        context['question'] = "Selecione um item do Quesito"                   
    return render(request, 'item/_itembynature_copy_item.html',context)

# ESSA FUNÇÃO APENAS CHECA QUE TIPO DE ITEM É, E VERIFICA SE ELE JÁ FOI CADASTRADO.
# SE LAUDO NÃO TEM QUESITO, ELE CRIA. SE TIVE QUESITO, VERIFICA QUAL O SEU TIPO E EDITA.
@login_required
def item_check_report(request, *args, **kwargs):
    report_id = kwargs.get('report',None)
    type_item_id = kwargs.get('type_item',None)
    marc = Item2.objects.filter(report_id=report_id, type_item_by_nature_of_action__type_item__id=type_item_id).exists()
    if marc:
        report = Report.objects.get(pk=report_id)
        #LEMBRAR: REGRA DE NEGÓCIO: SEMPRE VÃO SER 4 QUESITOS, DE TIPOS DIFERENTES, POR LAUDO. ENTÃO QUANDO VIER O QUESITO UM DOS 4 E ÚNICO. NUNCA VAI TER DOIS QUESITOS DA PROCURADORIA O DO JUIZ POR EXEMPLO.
        # POR ESSE MOTIVO QUE USO O GET AQUI. O CERTO SERIA O FILTER, MAS DEVIDO A REGRA DE NEGÓCIO ACIMA, POSSO USAR O GET
        get_item = Item2.objects.get(report=report_id, type_item_by_nature_of_action__type_item__id=type_item_id)      
        return redirect('url_item_edit', get_item.id)
    else:        
        return redirect('url_item_create2',report=report_id,type_item=type_item_id)


    item = get_object_or_404(Item, pk=pk)    
    if request.method == 'GET':        
        item.delete()
        messages.success(request, 'Ação concluída com sucesso.')
        return redirect('url_items_list')
    else:
        messages.warning(request, 'Ação não concluída.')
        return redirect('url_items_list')

@login_required
def item2_create(request, *args, **kwargs):
    print("Valor da request",request.POST)
    template_name = 'item2/form.html'
    report_id = kwargs.get('report',None)
    type_item_id = kwargs.get('type_item',None)  
    report = Report.objects.get(pk=report_id)
    type_item = TypeItem.objects.get(pk=type_item_id)   
    data = {}
    data['title'] = "Cadastro dos Quesitos"     
    data['report'] = report
    data['type_item'] = type_item      
    if request.method == 'POST':        
        form = Item2Form(request.POST)
        if form.is_valid():
            form.save()                        
            return redirect('url_report_edit',report.id)
        else:
            print("algo não está valido.")
    else:
        form = Item2Form()             
    
    data['form'] = form
    return render(request,template_name,data)

def copy_item_report(request, report, type_item):
    template_name = 'item2/copy_item_report.html'
    context = {}
    report = Report.objects.get(pk=report)
    type_item = TypeItem.objects.get(pk=type_item)
    #Verifica se foi dado um diagnóstico para este o laudo 
    cid_number = CidNumber.objects.filter(report_id=report.id, type_cid=True).exists()
    if cid_number:
        cid_number = CidNumber.objects.get(report_id=report.id, type_cid=True)
        #Verifica se tem modelo para este diagnóstico
        copy = TypeItemByNatureOfAction.objects.filter(nature_of_action=report.nature_of_action, cid_number=cid_number.category, type_item_id=type_item.id).exists()
        if copy:
            copy = TypeItemByNatureOfAction.objects.filter(nature_of_action=report.nature_of_action, cid_number=cid_number.category, type_item_id=type_item.id).order_by('id')
            context = {
                'type_item_by_nature_of_action': copy,
                'report': report,
                'type_item': type_item,
                'title': "Modelos de Perguntas",
            }   
            user = User.objects.get(username=request.user.username)
            if request.method == 'POST':
                marc = request.POST.get("quesitos","")
                print("MARC",marc)
                quesitos_ordenados = request.POST.getlist('quesitos')
                print("QUesitos ORdenados",quesitos_ordenados)
                if marc:                         
                    for item in quesitos_ordenados:
                        print("Item",item)
                        for cc in copy:                    
                            if int(item) == int(cc.id):  
                                if cc.answer:                      
                                    qq = Item2(report=report, type_item=type_item, question=cc.question, answer=cc.answer, answer_status=True, user_created=user, user_updated=user)
                                else:
                                    qq = Item2(report=report, type_item=type_item, question=cc.question, answer=cc.answer, answer_status=False, user_created=user, user_updated=user)
                                qq.save()
                return redirect('url_report_edit',report.id)
        else:            
            messages.warning(request, 'Não tem modelo para este diagnóstico.')
            return redirect('url_report_edit',report.id)
    else:
        messages.warning(request, 'Não existe diagnóstico para este Laudo.')
        return redirect('url_report_edit',report.id)
    
    return render(request,template_name,context)


@login_required
def items2_list(request, *args, **kwargs):    
    template_name = "item2/list.html"
    title = "Lista dos Quesitos"
    report_id = kwargs.get('report',None)
    type_item_id = kwargs.get('type_item',None)
    report = Report.objects.get(pk=report_id)
    type_item = TypeItem.objects.get(pk=type_item_id)     
    items2 = Item2.objects.filter(report_id=report_id, type_item_id=type_item_id).order_by('id')
    print("valores",items2)                         
    context = {
        'items': items2,
        'title': title,  
        'report': report,
        'type_item': type_item,                            
    }
    return render(request,template_name,context)

@login_required
def item2_detail(request, pk=None, *args, **kwargs):
    template_name = "item2/detail.html"
    item = get_object_or_404(Item,pk=pk)
    context = {
        'item2': item2
    }
    return render(request, template_name, context)

@login_required
def item2_edit(request, pk):    
    template_name='item2/form.html'
    request.session['tabs'] = "quesito-tab" # Isso é para quando redirecionar para o report, fica na tab certa
    title = "Editar a pergunta:"
    data = {}     
    item2 = get_object_or_404(Item2, pk=pk)  
    report = Report.objects.get(pk=item2.report_id)
    type_item = TypeItem.objects.get(pk=item2.type_item_id)
    data['report'] = report
    data['type_item'] = type_item 
    data['item2'] = item2   
    data['title'] = title     
    type_item_id = item2.type_item_id     
    user_created = item2.user_created # Esta linha faz com que o user_created não seja modificado, para mostrar quem criou esta pessoa
    form = Item2Form(request.POST or None, instance=item2)
    if form.is_valid():        
        item2 = form.save(commit=False)
        item2.user_created = user_created               
        if item2.answer != "":
            item2.answer_status = True
        else:
            item2.answer_status = False
        item2.save()           
        return redirect('url_items2_list', item2.report_id, item2.type_item_id)
    
    data['form'] = form
    return render(request,template_name,data)

def item2_answer(request, pk):
    template_name='item2/form.html'
    request.session['tabs'] = "quesito-tab" # Isso é para quando redirecionar para o report, fica na tab certa
    title = "Editar a resposta"
    data = {}     
    item2 = get_object_or_404(Item2, pk=pk)  
    report = Report.objects.get(pk=item2.report_id)
    type_item = TypeItem.objects.get(pk=item2.type_item_id)
    data['report'] = report
    data['type_item'] = type_item
    data['item2'] = item2 
    data['title'] = title 
    data['answer'] = "answer"# Marcador para saber se estou respondendo     
    type_item_id = item2.type_item_id     
    user_created = item2.user_created # Esta linha faz com que o user_created não seja modificado, para mostrar quem criou esta pessoa
    form = Item2Form(request.POST or None, instance=item2)
    if form.is_valid():        
        item2 = form.save(commit=False)
        item2.user_created = user_created               
        item2.save()           
        return redirect('url_items2_list', item2.report_id, item2.type_item_id)
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def item2_delete(request,pk):
    item2 = get_object_or_404(Item2, pk=pk)    
    if request.method == 'GET':        
        item2.delete()
        messages.success(request, 'Ação concluída com sucesso.')
        return redirect('url_items2_list', item2.report_id, item2.type_item_id)
    else:
        messages.warning(request, 'Ação não concluída.')
        return redirect('url_items2_list',item2.report_id, item2.type_item_id)

@login_required
def item2_delete_all(request):
    context = []
    context = request.GET.get("ff","")
    context = context.replace("[","").replace("]","").replace('\"','')
    context = context.split(",")
    context = [int(x) for x in context] 
    print("context",context)     
    if context:                
        b = Item2.objects.filter(id__in=context)
        b.delete()
        results = []         
        json = {}
        json['message'] = "Deletado"        
        results.append(json)
        return JsonResponse(results, safe=False)        
    else:
        json['message'] = "Não deletado"        
        results.append(json)
        return JsonResponse(results, safe=False)

##### FIM ITEM 2 #############################



######### ITEM ANSWER #################################

@login_required
def itemsanswer_create(request):
    template_name = 'itemsanswer/form.html'
    data = {}
    data['title'] = "Cadastro das Respostas dos Quesitos"     
    if request.method == 'POST':        
        form = ItemsAnswerForm(request.POST)
        if form.is_valid():
            form.save()                        
            return redirect('url_itemsanswers_list')
        else:
            print("algo não está valido.")
    else:
        form = ItemsAnswerForm()             
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def itemsanswers_list(request):
    template_name = "itemsanswer/list.html"
    title = "Lista das Resposta dos Quesitos"    
    itemsanswers = ItemsAnswer.objects.all()                        
    context = {
        'itemsanswers': itemsanswers,
        'title': title,                      
    }
    return render(request,template_name,context)

@login_required
def itemsanswer_detail(request, pk=None, *args, **kwargs):
    template_name = "itemsanswer/detail.html"
    itemsanswer = get_object_or_404(ItemsAnswer,pk=pk)
    context = {
        'itemsanswer': itemsanswer
    }
    return render(request, template_name, context)

@login_required
def itemsanswer_edit(request, pk):    
    template_name='itemsanswer/form.html'
    data = {}     
    itemsanswer = get_object_or_404(ItemsAnswer, pk=pk)        
    user_created = itemsanswer.user_created # Esta linha faz com que o user_created não seja modificado, para mostrar quem criou esta pessoa
    form = ItemsAnswerForm(request.POST or None, instance=itemsanswer)
    if form.is_valid():        
        itemsanswer = form.save(commit=False)
        itemsanswer.user_created = user_created               
        itemsanswer.save()
        return redirect('url_itemsanswer_detail',pk=pk)    
    
    data['form'] = form
    return render(request,template_name,data)

@login_required
def itemsanswer_delete(request,pk):
    itemsanswer = get_object_or_404(ItemsAnswer, pk=pk)    
    if request.method == 'GET':        
        itemsanswer.delete()
        messages.success(request, 'Ação concluída com sucesso.')
        return redirect('url_itemsanswers_list')
    else:
        messages.warning(request, 'Ação não concluída.')
        return redirect('url_itemsanswers_list')

######### FIM ITEM ANSWER #############################
